#!/usr/bin/env python3
"""
Script to create a comprehensive Word document for Deigratia Montessori School Management System documentation.
This script combines all documentation files into a single, professionally formatted DOCX file.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os

def create_documentation_docx():
    """Create a comprehensive Word document with all system documentation."""
    
    # Create a new document
    doc = Document()
    
    # Set up document styles
    setup_document_styles(doc)
    
    # Add title page
    add_title_page(doc)
    
    # Add table of contents placeholder
    add_table_of_contents(doc)
    
    # Add main documentation sections
    add_system_overview(doc)
    add_architecture_section(doc)
    add_user_roles_section(doc)
    add_core_features_section(doc)
    add_module_details_section(doc)
    add_authentication_section(doc)
    add_dashboard_section(doc)
    add_website_features_section(doc)
    add_technical_stack_section(doc)
    add_user_manual_section(doc)
    add_quick_reference_section(doc)
    add_troubleshooting_section(doc)
    
    # Save the document
    doc.save('Deigratia_School_Management_System_Documentation.docx')
    print("Documentation DOCX file created successfully!")
    print("File saved as: Deigratia_School_Management_System_Documentation.docx")

def setup_document_styles(doc):
    """Set up custom styles for the document."""
    
    # Title style
    title_style = doc.styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(28)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(24)
    
    # Subtitle style
    subtitle_style = doc.styles.add_style('Custom Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Calibri'
    subtitle_font.size = Pt(18)
    subtitle_font.color.rgb = RGBColor(0, 102, 153)  # Medium blue
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(18)
    
    # Heading 1 style (modify existing)
    heading1 = doc.styles['Heading 1']
    heading1.font.name = 'Calibri'
    heading1.font.size = Pt(20)
    heading1.font.color.rgb = RGBColor(0, 51, 102)
    heading1.paragraph_format.space_before = Pt(24)
    heading1.paragraph_format.space_after = Pt(12)
    
    # Heading 2 style (modify existing)
    heading2 = doc.styles['Heading 2']
    heading2.font.name = 'Calibri'
    heading2.font.size = Pt(16)
    heading2.font.color.rgb = RGBColor(0, 102, 153)
    heading2.paragraph_format.space_before = Pt(18)
    heading2.paragraph_format.space_after = Pt(6)
    
    # Heading 3 style (modify existing)
    heading3 = doc.styles['Heading 3']
    heading3.font.name = 'Calibri'
    heading3.font.size = Pt(14)
    heading3.font.color.rgb = RGBColor(51, 51, 51)
    heading3.paragraph_format.space_before = Pt(12)
    heading3.paragraph_format.space_after = Pt(6)
    
    # Normal style (modify existing)
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

def add_title_page(doc):
    """Add a professional title page."""
    
    # School name and title
    title = doc.add_paragraph('Deigratia Montessori School', style='Custom Title')
    
    subtitle = doc.add_paragraph('Management System Documentation', style='Custom Subtitle')
    
    # Add some space
    doc.add_paragraph()
    doc.add_paragraph()
    
    # System overview
    overview = doc.add_paragraph()
    overview.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = overview.add_run('Comprehensive System Documentation\n')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 51, 51)
    
    run2 = overview.add_run('Version 1.0 - Complete Feature Documentation')
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(102, 102, 102)
    
    # Add more space
    for _ in range(8):
        doc.add_paragraph()
    
    # School information
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info_run = info.add_run('Deigratia Montessori School\n')
    info_run.font.size = Pt(14)
    info_run.font.bold = True
    
    info_run2 = info.add_run('Oyibi, Greater Accra Region, Ghana\n')
    info_run2.font.size = Pt(12)
    
    info_run3 = info.add_run('https://deigratiams.edu.gh/')
    info_run3.font.size = Pt(12)
    info_run3.font.color.rgb = RGBColor(0, 102, 204)
    
    # Add page break
    doc.add_page_break()

def add_table_of_contents(doc):
    """Add table of contents."""
    
    doc.add_heading('Table of Contents', level=1)
    
    toc_items = [
        ('1. System Overview', '3'),
        ('2. Architecture', '5'),
        ('3. User Types and Roles', '7'),
        ('4. Core Features', '12'),
        ('5. Module Details', '15'),
        ('6. Authentication System', '25'),
        ('7. Dashboard System', '27'),
        ('8. Website Features', '30'),
        ('9. Technical Stack', '33'),
        ('10. User Manual', '35'),
        ('11. Quick Reference Guide', '50'),
        ('12. Troubleshooting', '55'),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item).font.size = Pt(12)
        
        # Add dots
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6))
        p.add_run('\t' + '.' * 50 + '\t' + page)
    
    doc.add_page_break()

def add_system_overview(doc):
    """Add system overview section."""
    
    doc.add_heading('1. System Overview', level=1)
    
    doc.add_paragraph(
        'The Deigratia Montessori School Management System (DGMS) is a comprehensive web-based platform '
        'designed specifically for Deigratia Montessori School in Oyibi, Greater Accra, Ghana. The system '
        'integrates both a public-facing website and a complete school management system (SMS) to handle '
        'all aspects of school operations.'
    )
    
    doc.add_heading('Key Characteristics', level=2)
    
    characteristics = [
        'Dual-purpose platform: Public website + Internal management system',
        'Role-based access control: 12 distinct user roles with specific permissions',
        'Montessori-focused: Tailored for Montessori educational methodology',
        'Cloud-integrated: Uses Cloudinary for media storage',
        'Mobile-responsive: Optimized for both desktop and mobile devices',
        'Production-ready: Currently deployed on Fly.io with PostgreSQL database'
    ]
    
    for char in characteristics:
        p = doc.add_paragraph(char, style='List Bullet')
    
    doc.add_heading('System Scope', level=2)
    
    doc.add_paragraph(
        'The system serves multiple stakeholders including administrators, teachers, students, parents, '
        'and various support staff members. It provides comprehensive functionality for academic management, '
        'financial operations, communication, and administrative tasks.'
    )

def add_architecture_section(doc):
    """Add architecture section."""
    
    doc.add_heading('2. Architecture', level=1)
    
    doc.add_heading('Project Structure', level=2)
    
    structure_text = """
DeigratiaMontessori/
├── ricas_school_manager/     # Main Django project
├── website/                  # Public website app
├── users/                    # User management & authentication
├── dashboard/                # Role-based dashboards
├── courses/                  # Course & class management
├── assignments/              # Assignment & grading system
├── attendance/               # Attendance tracking
├── fees/                     # Fee management
├── payroll/                  # Staff payroll system
├── communications/           # Internal messaging
├── appointments/             # Appointment booking
├── documents/                # Document management
├── templates/                # Shared templates
├── static/                   # Static files (CSS, JS, images)
└── media/                    # User uploaded files
    """
    
    code_para = doc.add_paragraph(structure_text)
    code_para.style = 'No Spacing'
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    doc.add_heading('Technology Stack', level=2)
    
    tech_items = [
        'Backend: Django 4.x (Python)',
        'Database: PostgreSQL (Production), SQLite (Development)',
        'Frontend: Bootstrap 4, HTML5, CSS3, JavaScript',
        'Media Storage: Cloudinary',
        'Deployment: Fly.io',
        'Email: SMTP configuration',
        'Rich Text: TinyMCE editor',
        'Authentication: Custom user model with role-based access'
    ]
    
    for item in tech_items:
        doc.add_paragraph(item, style='List Bullet')

def add_user_roles_section(doc):
    """Add user roles section."""
    
    doc.add_heading('3. User Types and Roles', level=1)
    
    roles = [
        {
            'name': 'ADMIN - System Administrator',
            'functions': [
                'Complete system access and control',
                'User management (create, edit, delete all user types)',
                'System configuration and settings',
                'Financial management oversight',
                'Report generation and analytics',
                'Bulk operations (CSV imports, mass communications)'
            ],
            'features': [
                'Advanced dashboard with system analytics',
                'User management with Excel export',
                'Bulk CSV upload for users',
                'System backup and restore',
                'Widget management for dashboards',
                'Complete fee and payroll oversight'
            ]
        },
        {
            'name': 'TEACHER - Teaching Staff',
            'functions': [
                'Class and subject management',
                'Assignment creation and grading',
                'Attendance tracking',
                'Student progress monitoring',
                'Parent communication'
            ],
            'features': [
                'Subject-specific dashboards',
                'Assignment management (create, grade, feedback)',
                'Attendance recording for assigned classes',
                'Grade book management',
                'Direct messaging with students and parents',
                'Course material upload and management'
            ]
        },
        {
            'name': 'STUDENT - Enrolled Students',
            'functions': [
                'Assignment submission and tracking',
                'Grade viewing',
                'Course material access',
                'Communication with teachers',
                'Schedule viewing'
            ],
            'features': [
                'Personal dashboard with academic overview',
                'Assignment submission portal',
                'Grade tracking and report cards',
                'Course material downloads',
                'Messaging system (teachers only)',
                'Attendance history viewing'
            ]
        },
        {
            'name': 'PARENT - Student Guardians',
            'functions': [
                'Child\'s academic progress monitoring',
                'Teacher communication',
                'Fee payment tracking',
                'Appointment scheduling',
                'Document management'
            ],
            'features': [
                'Multi-child management dashboard',
                'Real-time grade and attendance monitoring',
                'Direct messaging with teachers and administrators',
                'Fee payment history and outstanding balance tracking',
                'Appointment booking system',
                'Document upload for children'
            ]
        }
    ]
    
    for role in roles:
        doc.add_heading(role['name'], level=2)
        
        doc.add_heading('Primary Functions:', level=3)
        for func in role['functions']:
            doc.add_paragraph(func, style='List Bullet')
        
        doc.add_heading('Key Features:', level=3)
        for feature in role['features']:
            doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('Additional Staff Roles', level=2)
    
    staff_roles = [
        'ACCOUNTANT - Financial Management',
        'SECRETARY - Administrative Support',
        'RECEPTIONIST - Front Desk Operations',
        'SECURITY - Security Personnel',
        'JANITOR - Maintenance Staff',
        'COOK - Kitchen Staff',
        'CLEANER - Cleaning Staff',
        'STAFF - Other Support Staff'
    ]
    
    for role in staff_roles:
        doc.add_paragraph(role, style='List Bullet')

def add_core_features_section(doc):
    """Add core features section."""
    
    doc.add_heading('4. Core Features', level=1)
    
    features = [
        {
            'name': 'User Management System',
            'items': [
                'Custom User Model: Email-based authentication instead of username',
                'Role-based Access Control: 12 distinct roles with specific permissions',
                'Profile Management: Profile pictures, contact information, document uploads',
                'Bulk Operations: CSV import for mass user creation',
                'Flexible Authentication: Multiple login methods for different user types'
            ]
        },
        {
            'name': 'Academic Management',
            'items': [
                'Course Structure: Subjects, Classes, Sections',
                'Class Management: Teacher assignments, student enrollment',
                'Schedule Management: Timetable creation and management',
                'Academic Year/Term Management: Multi-term academic calendar'
            ]
        },
        {
            'name': 'Assignment & Grading System',
            'items': [
                'Assignment Types: Homework, Quiz, Test, Exam, Project',
                'Question Types: Multiple choice, True/False, Short answer, Long answer, File upload',
                'Grading System: Customizable grading scales and thresholds',
                'Report Cards: Automated report card generation',
                'Progress Tracking: Real-time academic progress monitoring'
            ]
        },
        {
            'name': 'Communication System',
            'items': [
                'Internal Messaging: Direct messaging between users',
                'Announcements: System-wide and targeted announcements',
                'Notifications: Real-time notification system',
                'Email Integration: Automated email notifications',
                'Role-based Communication: Restricted messaging based on user roles'
            ]
        }
    ]
    
    for feature in features:
        doc.add_heading(feature['name'], level=2)
        for item in feature['items']:
            doc.add_paragraph(item, style='List Bullet')

def add_module_details_section(doc):
    """Add module details section."""
    
    doc.add_heading('5. Module Details', level=1)
    
    modules = [
        {
            'name': 'Website Module (Public-Facing)',
            'purpose': 'Public website for Deigratia Montessori School',
            'features': [
                'Hero Slides: Dynamic homepage carousel with configurable buttons',
                'About Section: School information, staff profiles, Montessori methodology',
                'Academic Programs: Curriculum details, special programs, assessment methods',
                'Events: School events calendar and highlights',
                'News & Announcements: Public announcements and news',
                'Gallery: Photo galleries organized by categories',
                'Staff Directory: Public staff profiles and information',
                'Contact Forms: Inquiry forms and contact information',
                'Testimonials: Parent, student, and alumni testimonials',
                'FAQ Section: Frequently asked questions',
                'Career Portal: Job postings and application system'
            ]
        },
        {
            'name': 'Users Module',
            'purpose': 'Complete user management and authentication system',
            'features': [
                'Custom User Model: Extended Django user model with role-based fields',
                'Authentication Backends: Flexible student authentication system',
                'Profile Management: User profiles with additional information',
                'Role Management: Comprehensive role-based access control',
                'Bulk Import: CSV-based bulk user creation',
                'Email Integration: Welcome emails and notifications'
            ]
        },
        {
            'name': 'Dashboard Module',
            'purpose': 'Role-based dashboard system with customizable widgets',
            'features': [
                'Role-specific Dashboards: Tailored dashboards for each user role',
                'Widget System: Customizable dashboard widgets',
                'Theme Management: Multiple theme options (default, dark, teal)',
                'Preferences: User-specific dashboard preferences',
                'Analytics: Real-time system analytics and statistics',
                'Quick Actions: Role-based quick action menus'
            ]
        }
    ]
    
    for module in modules:
        doc.add_heading(module['name'], level=2)
        
        purpose_para = doc.add_paragraph()
        purpose_run = purpose_para.add_run('Purpose: ')
        purpose_run.bold = True
        purpose_para.add_run(module['purpose'])
        
        doc.add_heading('Key Features:', level=3)
        for feature in module['features']:
            doc.add_paragraph(feature, style='List Bullet')

def add_authentication_section(doc):
    """Add authentication section."""

    doc.add_heading('6. Authentication System', level=1)

    doc.add_heading('Login Methods', level=2)

    doc.add_paragraph('1. Email + Password: Standard login for all users')
    doc.add_paragraph('2. Student ID + PIN: Flexible student authentication')
    doc.add_paragraph('3. Flexible Student ID Formats: Supports various ID formats (case-insensitive, numeric-only, with/without prefixes)')

    doc.add_heading('Security Features', level=2)

    security_features = [
        'Role-based Access Control: Strict permission management',
        'Session Management: Secure session handling',
        'Password Security: Django\'s built-in password hashing',
        'Email Verification: User verification system',
        'Login Redirects: Role-based login redirects'
    ]

    for feature in security_features:
        doc.add_paragraph(feature, style='List Bullet')

def add_dashboard_section(doc):
    """Add dashboard section."""

    doc.add_heading('7. Dashboard System', level=1)

    doc.add_heading('Role-based Dashboards', level=2)

    dashboards = [
        {
            'name': 'Admin Dashboard',
            'features': [
                'System analytics and statistics',
                'User management quick actions',
                'Financial overview',
                'Recent activities',
                'System health monitoring'
            ]
        },
        {
            'name': 'Teacher Dashboard',
            'features': [
                'Class and subject overview',
                'Assignment management',
                'Student progress tracking',
                'Communication center',
                'Schedule display'
            ]
        },
        {
            'name': 'Student Dashboard',
            'features': [
                'Academic progress overview',
                'Assignment submissions',
                'Grade tracking',
                'Course materials access',
                'Communication center'
            ]
        },
        {
            'name': 'Parent Dashboard',
            'features': [
                'Children\'s academic progress',
                'Fee payment status',
                'Communication with teachers',
                'Appointment booking',
                'Document management'
            ]
        }
    ]

    for dashboard in dashboards:
        doc.add_heading(dashboard['name'], level=3)
        for feature in dashboard['features']:
            doc.add_paragraph(feature, style='List Bullet')

def add_website_features_section(doc):
    """Add website features section."""

    doc.add_heading('8. Website Features', level=1)

    doc.add_heading('Public Website Components', level=2)

    components = [
        {
            'name': 'Homepage',
            'features': [
                'Dynamic hero slides with configurable content',
                'Featured announcements and news',
                'Upcoming events display',
                'Parent/student testimonials',
                'Photo gallery carousel',
                'Quick access to key information'
            ]
        },
        {
            'name': 'About Section',
            'features': [
                'School history and mission',
                'Montessori methodology explanation',
                'Staff profiles and qualifications',
                'School facilities and campus tour',
                'Educational philosophy'
            ]
        },
        {
            'name': 'Academic Programs',
            'features': [
                'Curriculum details and approach',
                'Special programs and activities',
                'Assessment methods',
                'Academic calendar',
                'Admission requirements'
            ]
        }
    ]

    for component in components:
        doc.add_heading(component['name'], level=3)
        for feature in component['features']:
            doc.add_paragraph(feature, style='List Bullet')

def add_technical_stack_section(doc):
    """Add technical stack section."""

    doc.add_heading('9. Technical Stack', level=1)

    doc.add_heading('Database Design', level=2)

    db_features = [
        'PostgreSQL: Production database',
        'Relational Structure: Well-normalized database design',
        'Foreign Key Relationships: Proper data relationships',
        'Indexing: Optimized database queries',
        'Data Integrity: Constraints and validation'
    ]

    for feature in db_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Security Measures', level=2)

    security_measures = [
        'HTTPS: Secure data transmission',
        'CSRF Protection: Cross-site request forgery protection',
        'SQL Injection Prevention: Parameterized queries',
        'XSS Protection: Cross-site scripting prevention',
        'Role-based Access: Strict permission controls'
    ]

    for measure in security_measures:
        doc.add_paragraph(measure, style='List Bullet')

def add_user_manual_section(doc):
    """Add user manual section."""

    doc.add_heading('10. User Manual', level=1)

    doc.add_heading('Getting Started', level=2)

    doc.add_paragraph('System Access:')
    doc.add_paragraph('• Website: https://deigratiams.edu.gh/')
    doc.add_paragraph('• Login Page: Click "Login" on the website or go to /users/login/')
    doc.add_paragraph('• Support: Contact school administration for login credentials')

    doc.add_heading('Login Methods', level=3)

    doc.add_paragraph('For Students:')
    doc.add_paragraph('1. Student ID + PIN: Use your student ID and 5-digit PIN')
    doc.add_paragraph('2. Email + Password: Use your email address and password')
    doc.add_paragraph('3. Flexible ID Format: Student ID can be entered in various formats')

    doc.add_paragraph('For All Other Users:')
    doc.add_paragraph('• Email + Password: Use your registered email address and password')

def add_quick_reference_section(doc):
    """Add quick reference section."""

    doc.add_heading('11. Quick Reference Guide', level=1)

    doc.add_heading('System Overview', level=2)

    overview_items = [
        'School: Deigratia Montessori School, Oyibi, Greater Accra, Ghana',
        'Domain: https://deigratiams.edu.gh/',
        'Platform: Django-based web application',
        'Deployment: Fly.io with PostgreSQL database',
        'Storage: Cloudinary for media files'
    ]

    for item in overview_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Key URLs & Access Points', level=2)

    urls = [
        'Homepage: /',
        'Login: /users/login/',
        'Admin Panel: /my-admin/',
        'Dashboard: /dashboard/',
        'Courses: /courses/',
        'Assignments: /assignments/',
        'Communications: /communications/'
    ]

    for url in urls:
        doc.add_paragraph(url, style='List Bullet')

def add_troubleshooting_section(doc):
    """Add troubleshooting section."""

    doc.add_heading('12. Troubleshooting', level=1)

    doc.add_heading('Login Issues', level=2)

    doc.add_paragraph('Problem: Cannot log in')
    doc.add_paragraph('Solutions:')
    solutions = [
        'Check your email address and password',
        'For students: Try different student ID formats',
        'Clear browser cache and cookies',
        'Contact school administration for password reset'
    ]

    for solution in solutions:
        doc.add_paragraph(solution, style='List Bullet')

    doc.add_heading('Performance Issues', level=2)

    doc.add_paragraph('Problem: System is slow')
    doc.add_paragraph('Solutions:')
    perf_solutions = [
        'Check your internet connection',
        'Clear browser cache',
        'Try a different browser',
        'Contact IT support if problem persists'
    ]

    for solution in perf_solutions:
        doc.add_paragraph(solution, style='List Bullet')

    doc.add_heading('Getting Help', level=2)

    help_items = [
        'Technical Issues: Contact school IT support',
        'Account Issues: Contact school administration',
        'Academic Questions: Contact your teachers directly',
        'Fee Questions: Contact the school accountant'
    ]

    for item in help_items:
        doc.add_paragraph(item, style='List Bullet')

if __name__ == '__main__':
    create_documentation_docx()
