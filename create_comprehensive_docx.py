#!/usr/bin/env python3
"""
Enhanced script to create a comprehensive Word document for Deigratia Montessori School Management System documentation.
This script creates a more detailed and professionally formatted document with tables, better styling, and complete content.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import os

def create_comprehensive_documentation():
    """Create a comprehensive Word document with all system documentation."""
    
    # Create a new document
    doc = Document()
    
    # Set up document styles
    setup_enhanced_styles(doc)
    
    # Add title page
    add_enhanced_title_page(doc)
    
    # Add executive summary
    add_executive_summary(doc)
    
    # Add table of contents
    add_enhanced_toc(doc)
    
    # Add all sections
    add_complete_system_overview(doc)
    add_complete_user_roles(doc)
    add_complete_features(doc)
    add_complete_modules(doc)
    add_complete_technical_details(doc)
    add_complete_user_guides(doc)
    add_deployment_info(doc)
    add_appendices(doc)
    
    # Save the document
    doc.save('Deigratia_School_Management_System_Complete_Documentation.docx')
    print("Comprehensive documentation DOCX file created successfully!")
    print("File saved as: Deigratia_School_Management_System_Complete_Documentation.docx")

def setup_enhanced_styles(doc):
    """Set up enhanced custom styles for the document."""
    
    # Title style
    title_style = doc.styles.add_style('Enhanced Title', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(32)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(30)
    
    # Subtitle style
    subtitle_style = doc.styles.add_style('Enhanced Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Calibri'
    subtitle_font.size = Pt(20)
    subtitle_font.color.rgb = RGBColor(0, 102, 153)  # Medium blue
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(20)
    
    # Section header style
    section_style = doc.styles.add_style('Section Header', WD_STYLE_TYPE.PARAGRAPH)
    section_font = section_style.font
    section_font.name = 'Calibri'
    section_font.size = Pt(16)
    section_font.bold = True
    section_font.color.rgb = RGBColor(51, 51, 51)
    section_style.paragraph_format.space_before = Pt(12)
    section_style.paragraph_format.space_after = Pt(6)
    
    # Highlight style
    highlight_style = doc.styles.add_style('Highlight', WD_STYLE_TYPE.PARAGRAPH)
    highlight_font = highlight_style.font
    highlight_font.name = 'Calibri'
    highlight_font.size = Pt(12)
    highlight_font.bold = True
    highlight_font.color.rgb = RGBColor(0, 102, 153)
    highlight_style.paragraph_format.space_after = Pt(6)
    
    # Code style
    code_style = doc.styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
    code_font = code_style.font
    code_font.name = 'Courier New'
    code_font.size = Pt(9)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.space_after = Pt(6)

def add_enhanced_title_page(doc):
    """Add an enhanced professional title page."""
    
    # Add some space at the top
    for _ in range(3):
        doc.add_paragraph()
    
    # Main title
    title = doc.add_paragraph('Deigratia Montessori School', style='Enhanced Title')
    
    # Subtitle
    subtitle = doc.add_paragraph('Management System', style='Enhanced Subtitle')
    
    # Documentation type
    doc_type = doc.add_paragraph()
    doc_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_type.add_run('Complete System Documentation')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Add space
    for _ in range(4):
        doc.add_paragraph()
    
    # System information box
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create a bordered text box effect
    info_text = """
    🏫 SCHOOL INFORMATION
    
    Deigratia Montessori School
    Oyibi, Greater Accra Region, Ghana
    
    🌐 Website: https://deigratiams.edu.gh/
    📧 Email: info@deigratiams.edu.gh
    
    📋 SYSTEM DETAILS
    
    Platform: Django Web Application
    Database: PostgreSQL
    Deployment: Fly.io Cloud Platform
    Version: 1.0 (Production Ready)
    """
    
    info_run = info_para.add_run(info_text)
    info_run.font.size = Pt(12)
    info_run.font.name = 'Calibri'
    
    # Add more space
    for _ in range(6):
        doc.add_paragraph()
    
    # Footer information
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_run = footer.add_run('Comprehensive Documentation\nIncluding User Manuals, Technical Specifications, and System Features')
    footer_run.font.size = Pt(11)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Add page break
    doc.add_page_break()

def add_executive_summary(doc):
    """Add executive summary section."""
    
    doc.add_heading('Executive Summary', level=1)
    
    doc.add_paragraph(
        'The Deigratia Montessori School Management System (DGMS) represents a comprehensive digital '
        'transformation solution specifically designed for Deigratia Montessori School in Oyibi, Ghana. '
        'This state-of-the-art web-based platform seamlessly integrates academic management, financial '
        'operations, communication systems, and administrative functions into a unified, user-friendly interface.'
    )
    
    doc.add_heading('Key Achievements', level=2)
    
    achievements = [
        '✅ Successfully deployed and operational at https://deigratiams.edu.gh/',
        '✅ Serving 12 distinct user roles with customized access and functionality',
        '✅ Managing complete academic lifecycle from enrollment to graduation',
        '✅ Processing financial transactions and payroll operations',
        '✅ Facilitating seamless communication between all stakeholders',
        '✅ Providing real-time analytics and comprehensive reporting',
        '✅ Ensuring mobile-responsive access for all users'
    ]
    
    for achievement in achievements:
        doc.add_paragraph(achievement, style='List Bullet')
    
    doc.add_heading('System Impact', level=2)
    
    doc.add_paragraph(
        'The implementation of DGMS has revolutionized school operations by automating manual processes, '
        'improving communication efficiency, enhancing academic tracking, and providing transparent '
        'financial management. The system serves as a central hub for all school activities, ensuring '
        'consistency, accuracy, and accessibility of information for all stakeholders.'
    )
    
    doc.add_page_break()

def add_enhanced_toc(doc):
    """Add enhanced table of contents."""
    
    doc.add_heading('Table of Contents', level=1)
    
    toc_sections = [
        ('Executive Summary', '2'),
        ('1. System Overview', '4'),
        ('   1.1 Introduction', '4'),
        ('   1.2 System Architecture', '5'),
        ('   1.3 Key Features', '6'),
        ('2. User Roles and Permissions', '8'),
        ('   2.1 Administrative Roles', '8'),
        ('   2.2 Academic Roles', '10'),
        ('   2.3 Support Staff Roles', '12'),
        ('3. Core System Features', '14'),
        ('   3.1 User Management', '14'),
        ('   3.2 Academic Management', '16'),
        ('   3.3 Communication System', '18'),
        ('   3.4 Financial Management', '20'),
        ('4. Module Documentation', '22'),
        ('   4.1 Website Module', '22'),
        ('   4.2 Users Module', '24'),
        ('   4.3 Dashboard Module', '26'),
        ('   4.4 Academic Modules', '28'),
        ('   4.5 Administrative Modules', '32'),
        ('5. Technical Specifications', '36'),
        ('   5.1 Technology Stack', '36'),
        ('   5.2 Security Features', '38'),
        ('   5.3 Performance Optimization', '40'),
        ('6. User Guides', '42'),
        ('   6.1 Administrator Guide', '42'),
        ('   6.2 Teacher Guide', '46'),
        ('   6.3 Student Guide', '50'),
        ('   6.4 Parent Guide', '54'),
        ('7. Deployment and Maintenance', '58'),
        ('   7.1 Production Environment', '58'),
        ('   7.2 Backup and Recovery', '60'),
        ('   7.3 Monitoring and Support', '62'),
        ('8. Appendices', '64'),
        ('   8.1 API Documentation', '64'),
        ('   8.2 Database Schema', '66'),
        ('   8.3 Troubleshooting Guide', '68'),
    ]
    
    for section, page in toc_sections:
        p = doc.add_paragraph()
        p.add_run(section).font.size = Pt(11)
        
        # Add dots and page number
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6))
        dots = '.' * (80 - len(section))
        p.add_run(f'\t{dots}\t{page}')
    
    doc.add_page_break()

def add_complete_system_overview(doc):
    """Add complete system overview section."""
    
    doc.add_heading('1. System Overview', level=1)
    
    doc.add_heading('1.1 Introduction', level=2)
    
    doc.add_paragraph(
        'The Deigratia Montessori School Management System (DGMS) is a cutting-edge, web-based platform '
        'that serves as the digital backbone for Deigratia Montessori School. Located in Oyibi, Greater '
        'Accra Region, Ghana, this institution has embraced modern technology to enhance educational '
        'delivery and administrative efficiency.'
    )
    
    # Create a features table
    doc.add_heading('System Capabilities Overview', level=3)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Features'
    hdr_cells[2].text = 'Users Served'
    
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add data rows
    capabilities = [
        ('Academic Management', 'Courses, Assignments, Grading, Attendance', 'Teachers, Students, Parents'),
        ('User Management', 'Role-based Access, Bulk Import, Profiles', 'Administrators'),
        ('Communication', 'Messaging, Announcements, Notifications', 'All Users'),
        ('Financial Management', 'Fees, Payments, Payroll, Receipts', 'Accountants, Parents'),
        ('Website Management', 'Public Content, SEO, Gallery, Events', 'Public, Administrators'),
        ('Document Management', 'Secure Storage, Approval Workflow', 'All Users'),
        ('Appointment System', 'Booking, Scheduling, Reminders', 'Parents, Staff'),
        ('Analytics & Reporting', 'Performance Metrics, Usage Statistics', 'Administrators, Teachers')
    ]
    
    for category, features, users in capabilities:
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = features
        row_cells[2].text = users
    
    doc.add_heading('1.2 System Architecture', level=2)
    
    doc.add_paragraph(
        'DGMS follows a modular, scalable architecture built on the Django web framework. The system '
        'is designed with separation of concerns, ensuring maintainability, security, and performance.'
    )
    
    # Architecture components
    arch_components = [
        'Frontend Layer: Bootstrap 4, HTML5, CSS3, JavaScript',
        'Application Layer: Django 4.x with Python',
        'Database Layer: PostgreSQL with optimized queries',
        'Storage Layer: Cloudinary for media files',
        'Security Layer: Role-based access control, HTTPS, CSRF protection',
        'Integration Layer: Email services, cloud storage APIs'
    ]
    
    for component in arch_components:
        doc.add_paragraph(component, style='List Bullet')

def add_complete_user_roles(doc):
    """Add complete user roles section with detailed tables."""
    
    doc.add_heading('2. User Roles and Permissions', level=1)
    
    doc.add_heading('2.1 Administrative Roles', level=2)
    
    # Admin role detailed table
    doc.add_heading('System Administrator (ADMIN)', level=3)
    
    admin_table = doc.add_table(rows=1, cols=2)
    admin_table.style = 'Table Grid'
    
    # Header
    admin_hdr = admin_table.rows[0].cells
    admin_hdr[0].text = 'Capability'
    admin_hdr[1].text = 'Description'
    
    admin_capabilities = [
        ('User Management', 'Create, edit, delete users; bulk import via CSV; role assignment'),
        ('System Configuration', 'School settings, academic terms, grading scales, email configuration'),
        ('Financial Oversight', 'Complete fee management, payroll processing, financial reporting'),
        ('Content Management', 'Website content, announcements, events, gallery management'),
        ('Analytics & Reporting', 'System usage statistics, academic performance reports, financial analytics'),
        ('Bulk Operations', 'Mass communications, bulk user operations, system maintenance'),
        ('Security Management', 'User permissions, system security settings, audit logs')
    ]
    
    for capability, description in admin_capabilities:
        row_cells = admin_table.add_row().cells
        row_cells[0].text = capability
        row_cells[1].text = description
    
    doc.add_heading('2.2 Academic Roles', level=2)
    
    # Teacher role
    doc.add_heading('Teacher (TEACHER)', level=3)
    
    teacher_features = [
        'Class and subject management with student enrollment',
        'Assignment creation with multiple question types',
        'Automated and manual grading systems',
        'Daily attendance tracking and reporting',
        'Grade book management with progress tracking',
        'Direct communication with students and parents',
        'Course material upload and organization',
        'Student performance analytics and reporting'
    ]
    
    for feature in teacher_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # Student role
    doc.add_heading('Student (STUDENT)', level=3)
    
    student_features = [
        'Personal academic dashboard with progress overview',
        'Assignment submission portal with file upload',
        'Real-time grade tracking and report card access',
        'Course material downloads and resource access',
        'Direct messaging with teachers',
        'Attendance history and schedule viewing',
        'Academic calendar and event notifications'
    ]
    
    for feature in student_features:
        doc.add_paragraph(feature, style='List Bullet')

def add_complete_features(doc):
    """Add complete features section."""
    
    doc.add_heading('3. Core System Features', level=1)
    
    doc.add_heading('3.1 User Management System', level=2)
    
    doc.add_paragraph('Advanced User Management Capabilities:', style='Highlight')
    
    user_mgmt_features = [
        'Custom User Model: Email-based authentication with role-specific fields',
        'Flexible Authentication: Multiple login methods including student ID + PIN',
        'Role-based Access Control: 12 distinct roles with granular permissions',
        'Bulk Import System: CSV-based user creation with progress tracking',
        'Profile Management: Comprehensive user profiles with document uploads',
        'Welcome Email Automation: Automatic credential delivery to new users',
        'User Analytics: Registration trends and user activity monitoring'
    ]
    
    for feature in user_mgmt_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('3.2 Academic Management System', level=2)
    
    doc.add_paragraph('Comprehensive Academic Operations:', style='Highlight')
    
    academic_features = [
        'Course Structure: Hierarchical organization of subjects, classes, and sections',
        'Assignment System: Multiple assignment types with varied question formats',
        'Grading Engine: Automated grading for MCQs with manual override capabilities',
        'Report Cards: Automated generation with customizable templates',
        'Attendance Tracking: Daily attendance with multiple status options',
        'Progress Monitoring: Real-time academic progress tracking and analytics',
        'Schedule Management: Comprehensive timetable creation and management'
    ]
    
    for feature in academic_features:
        doc.add_paragraph(feature, style='List Bullet')

def add_complete_modules(doc):
    """Add complete modules section."""

    doc.add_heading('4. Module Documentation', level=1)

    doc.add_heading('4.1 Website Module (Public Interface)', level=2)

    doc.add_paragraph('Public-Facing Features:', style='Highlight')

    website_features = [
        'Dynamic Hero Slides: Configurable homepage carousel with custom buttons',
        'About Section: School information, staff profiles, Montessori methodology',
        'Academic Programs: Curriculum details, special programs, assessment methods',
        'Events Calendar: School events with registration and management',
        'News & Announcements: Public news system with categorization',
        'Photo Gallery: Organized image galleries with categories',
        'Staff Directory: Public staff profiles with contact information',
        'Contact System: Multiple contact forms and inquiry management',
        'Testimonials: Parent, student, and alumni testimonials',
        'SEO Optimization: Search engine optimization features'
    ]

    for feature in website_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('4.2 Dashboard Module (User Interface)', level=2)

    doc.add_paragraph('Role-Specific Dashboard Features:', style='Highlight')

    dashboard_features = [
        'Customizable Widgets: User-configurable dashboard components',
        'Theme System: Multiple color schemes (Default, Dark, Teal)',
        'Real-time Analytics: Live system statistics and performance metrics',
        'Quick Actions: Role-based quick access menus',
        'Notification Center: Centralized notification management',
        'Responsive Design: Mobile-optimized dashboard layouts',
        'User Preferences: Personalized dashboard settings'
    ]

    for feature in dashboard_features:
        doc.add_paragraph(feature, style='List Bullet')

def add_complete_technical_details(doc):
    """Add complete technical details section."""

    doc.add_heading('5. Technical Specifications', level=1)

    doc.add_heading('5.1 Technology Stack', level=2)

    # Create technology stack table
    tech_table = doc.add_table(rows=1, cols=3)
    tech_table.style = 'Table Grid'

    tech_hdr = tech_table.rows[0].cells
    tech_hdr[0].text = 'Layer'
    tech_hdr[1].text = 'Technology'
    tech_hdr[2].text = 'Purpose'

    tech_stack = [
        ('Backend Framework', 'Django 4.x', 'Web application framework with ORM'),
        ('Database', 'PostgreSQL', 'Production database with ACID compliance'),
        ('Frontend', 'Bootstrap 4 + JavaScript', 'Responsive UI framework'),
        ('Media Storage', 'Cloudinary', 'Cloud-based media storage and optimization'),
        ('Deployment', 'Fly.io', 'Cloud application platform'),
        ('Email Service', 'SMTP', 'Email delivery and notifications'),
        ('Rich Text Editor', 'TinyMCE', 'WYSIWYG content editing'),
        ('Authentication', 'Django Auth + Custom', 'User authentication and authorization')
    ]

    for layer, tech, purpose in tech_stack:
        row_cells = tech_table.add_row().cells
        row_cells[0].text = layer
        row_cells[1].text = tech
        row_cells[2].text = purpose

    doc.add_heading('5.2 Security Implementation', level=2)

    security_measures = [
        'HTTPS Encryption: All data transmission secured with SSL/TLS',
        'CSRF Protection: Cross-site request forgery prevention',
        'SQL Injection Prevention: Parameterized queries and ORM protection',
        'XSS Protection: Cross-site scripting prevention measures',
        'Role-based Access Control: Granular permission system',
        'Session Security: Secure session management and timeout',
        'File Upload Security: File type validation and secure storage',
        'Password Security: Strong hashing with Django\'s PBKDF2 algorithm'
    ]

    for measure in security_measures:
        doc.add_paragraph(measure, style='List Bullet')

def add_complete_user_guides(doc):
    """Add complete user guides section."""

    doc.add_heading('6. User Guides', level=1)

    doc.add_heading('6.1 Administrator Guide', level=2)

    doc.add_paragraph('Getting Started as Administrator:', style='Highlight')

    admin_steps = [
        '1. Access the admin panel at /my-admin/ using your administrator credentials',
        '2. Configure school settings including basic information and academic terms',
        '3. Set up user roles and create initial user accounts',
        '4. Configure fee structure and payment methods',
        '5. Set up email settings for automated notifications',
        '6. Customize website content and public information',
        '7. Create academic calendar and important events'
    ]

    for step in admin_steps:
        doc.add_paragraph(step)

    doc.add_heading('Key Administrative Tasks:', level=3)

    admin_tasks = [
        'User Management: Create, edit, and manage all user accounts',
        'Bulk Operations: Import users via CSV and send mass communications',
        'Financial Oversight: Monitor fee payments and process payroll',
        'System Monitoring: Review system analytics and performance metrics',
        'Content Management: Update website content and school information',
        'Report Generation: Create and export various system reports'
    ]

    for task in admin_tasks:
        doc.add_paragraph(task, style='List Bullet')

    doc.add_heading('6.2 Teacher Guide', level=2)

    doc.add_paragraph('Teacher Dashboard Overview:', style='Highlight')

    teacher_guide = [
        'Access your teacher dashboard to view assigned classes and subjects',
        'Create assignments using the assignment builder with various question types',
        'Take daily attendance for your classes with status tracking',
        'Grade student submissions and provide detailed feedback',
        'Upload course materials and resources for student access',
        'Communicate with students and parents through the messaging system',
        'Generate progress reports and academic analytics'
    ]

    for guide in teacher_guide:
        doc.add_paragraph(guide, style='List Bullet')

def add_deployment_info(doc):
    """Add deployment information section."""

    doc.add_heading('7. Deployment and Maintenance', level=1)

    doc.add_heading('7.1 Production Environment', level=2)

    # Production details table
    prod_table = doc.add_table(rows=1, cols=2)
    prod_table.style = 'Table Grid'

    prod_hdr = prod_table.rows[0].cells
    prod_hdr[0].text = 'Component'
    prod_hdr[1].text = 'Configuration'

    prod_config = [
        ('Domain', 'https://deigratiams.edu.gh/'),
        ('Hosting Platform', 'Fly.io Cloud Platform'),
        ('Database', 'PostgreSQL with automated backups'),
        ('SSL Certificate', 'Automatic HTTPS with Let\'s Encrypt'),
        ('Media Storage', 'Cloudinary CDN'),
        ('Email Service', 'SMTP with authentication'),
        ('Monitoring', 'Application performance monitoring'),
        ('Backup Schedule', 'Daily automated backups')
    ]

    for component, config in prod_config:
        row_cells = prod_table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = config

    doc.add_heading('7.2 System Maintenance', level=2)

    maintenance_tasks = [
        'Regular database backups and integrity checks',
        'System updates and security patches',
        'Performance monitoring and optimization',
        'User account management and cleanup',
        'Content updates and website maintenance',
        'Email system monitoring and configuration',
        'Storage usage monitoring and cleanup'
    ]

    for task in maintenance_tasks:
        doc.add_paragraph(task, style='List Bullet')

def add_appendices(doc):
    """Add appendices section."""

    doc.add_heading('8. Appendices', level=1)

    doc.add_heading('8.1 System URLs Reference', level=2)

    # URLs table
    url_table = doc.add_table(rows=1, cols=3)
    url_table.style = 'Table Grid'

    url_hdr = url_table.rows[0].cells
    url_hdr[0].text = 'Function'
    url_hdr[1].text = 'URL Path'
    url_hdr[2].text = 'Access Level'

    system_urls = [
        ('Homepage', '/', 'Public'),
        ('User Login', '/users/login/', 'All Users'),
        ('Admin Panel', '/my-admin/', 'Administrators'),
        ('Dashboard', '/dashboard/', 'Authenticated Users'),
        ('Courses', '/courses/', 'Teachers, Students'),
        ('Assignments', '/assignments/', 'Teachers, Students'),
        ('Communications', '/communications/', 'All Users'),
        ('Fees', '/fees/', 'Accountants, Parents'),
        ('Appointments', '/appointments/', 'Parents, Staff')
    ]

    for function, url, access in system_urls:
        row_cells = url_table.add_row().cells
        row_cells[0].text = function
        row_cells[1].text = url
        row_cells[2].text = access

    doc.add_heading('8.2 Troubleshooting Quick Reference', level=2)

    troubleshooting = [
        'Login Issues: Check credentials, clear browser cache, contact admin',
        'Performance Issues: Check internet connection, try different browser',
        'File Upload Problems: Verify file size and format requirements',
        'Email Not Received: Check spam folder, verify email address',
        'Permission Denied: Contact administrator to verify user role',
        'System Errors: Contact IT support with error details'
    ]

    for issue in troubleshooting:
        doc.add_paragraph(issue, style='List Bullet')

    doc.add_heading('8.3 Contact Information', level=2)

    contact_info = [
        'Technical Support: Contact school IT department',
        'Account Issues: Contact school administration',
        'Academic Questions: Contact respective teachers',
        'Financial Queries: Contact school accountant',
        'System Feedback: Submit through admin panel'
    ]

    for contact in contact_info:
        doc.add_paragraph(contact, style='List Bullet')

    # Add final page with document information
    doc.add_page_break()

    doc.add_heading('Document Information', level=1)

    doc_info = doc.add_paragraph()
    doc_info.add_run('Document Title: ').bold = True
    doc_info.add_run('Deigratia Montessori School Management System - Complete Documentation\n')

    doc_info.add_run('Version: ').bold = True
    doc_info.add_run('1.0\n')

    doc_info.add_run('Last Updated: ').bold = True
    doc_info.add_run('December 2024\n')

    doc_info.add_run('Prepared for: ').bold = True
    doc_info.add_run('Deigratia Montessori School\n')

    doc_info.add_run('System Status: ').bold = True
    doc_info.add_run('Production Ready and Deployed\n')

if __name__ == '__main__':
    create_comprehensive_documentation()
